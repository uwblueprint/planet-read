import os
from base64 import b64encode
from datetime import datetime, timedelta

import docx
from docx.enum.style import WD_STYLE_TYPE
from flask import current_app
from sqlalchemy import or_
from sqlalchemy.orm import aliased
from werkzeug.utils import secure_filename

from ...graphql.types.file_type import FileDTO
from ...graphql.types.story_type import (
    StageEnum,
    StoryTranslationContentResponseDTO,
    StoryTranslationStatisticsResponseDTO,
    StoryTranslationUpdateStatusResponseDTO,
)
from ...models import db
from ...models.comment import Comment
from ...models.file import File
from ...models.story import Story
from ...models.story_all import StoryAll
from ...models.story_content import StoryContent
from ...models.story_translation import StoryTranslation
from ...models.story_translation_all import StoryTranslationAll
from ...models.story_translation_content import StoryTranslationContent
from ...models.story_translation_content_status import StoryTranslationContentStatus
from ...models.user import User
from ..interfaces.story_service import IStoryService


class StoryService(IStoryService):
    def __init__(self, logger=current_app.logger):
        self.logger = logger

    def get_stories(self, story_title=None, start_date=None, end_date=None):
        try:
            # Story is a SQLAlchemy model, we can use convenient methods provided
            # by SQLAlchemy like query.all() to query the data
            filters = []
            if (end_date and start_date) and end_date < start_date:
                raise Exception("Invalid filter: end_date cannot be before start_date.")
            if story_title is not None:
                filters.append(StoryAll.title.like(f"%{story_title}%"))
            if start_date is not None:
                filters.append(StoryAll.date_uploaded >= start_date)
            if end_date is not None:
                filters.append(StoryAll.date_uploaded <= end_date)
            return [
                story.to_dict(include_relationships=True)
                for story in StoryAll.query.filter(*filters).all()
                if not story.is_deleted
            ]
        except Exception as error:
            self.logger.error(error)
            raise error

    def get_story(self, id):
        # get queries by the primary key, which is id for the Story table
        # note: include_relationships=True will only include StoryContents if StoryAll is called
        # There are no security/permission issues by allowing users to view a deleted story,
        # so StoryAll can be safely used here
        story = StoryAll.query.get(id)
        if story is None:
            self.logger.error("Invalid id")
            raise Exception("Invalid id")
        return story.to_dict(include_relationships=True)

    def create_story(self, story, content):
        # create story
        try:
            new_story = StoryAll(**story.__dict__)
            new_story.date_uploaded = datetime.utcnow()
        except Exception as error:
            self.logger.error(str(error))
            raise error

        db.session.add(new_story)
        db.session.commit()

        # insert contents into story_contents
        # TODO use batch inserts here instead of iterating over the for loop?
        try:
            for i, line in enumerate(content):
                new_content = {
                    "story_id": new_story.id,
                    "line_index": i,
                    "content": line,
                    "is_deleted": False,
                }
                db.session.add(StoryContent(**new_content))
        except Exception as error:
            self.logger.error(str(error))
            raise error

        db.session.commit()
        db.session.refresh(new_story)

        return new_story

    def import_story(self, details, file):
        try:
            story_contents = self.process_story(file["path"])
            return self.create_story(details, story_contents)

        except Exception as error:
            self.logger.error(str(error))
            raise error

    def export_story_translation(self, id):
        try:
            story_details = self.get_story_translation(id=id)
            doc = docx.Document()
            doc.add_heading(story_details["title"], 0)
            styles = doc.styles
            bold_style = styles.add_style("Bold", WD_STYLE_TYPE.PARAGRAPH)
            bold_style.font.bold = True
            doc.add_paragraph(
                f"Translator: {story_details['translator_name']}", style="Bold"
            )
            doc.add_paragraph(
                f"Reviewer: {story_details['reviewer_name']}", style="Bold"
            )
            doc.add_paragraph(f"Language: {story_details['language']}", style="Bold")
            doc.add_paragraph(f"Level: {story_details['level']}", style="Bold")

            for line in story_details["translation_contents"]:
                doc.add_paragraph(line["translation_content"])

            upload_folder_path = os.getenv("UPLOAD_PATH")

            upload_path = os.path.join(
                upload_folder_path,
                secure_filename("downloadable-story-translation.docx"),
            )
            doc.save(upload_path)
            file_dto = FileDTO(path=upload_path)
            self.logger.info(file_dto)
            new_file = File(**file_dto.__dict__)
            return new_file.to_dict()
        except Exception as error:
            self.logger.error(str(error))
            raise error

    def create_translation(self, translation):
        try:
            new_story_translation = StoryTranslation(**translation.__dict__)
            story_translations_translating = (
                self._get_story_translations_user_translating_query(
                    new_story_translation.translator_id, isTranslator=True
                ).all()
            )
            languages_currently_translating = self._get_story_translation_languages(
                story_translations_translating
            )
            if new_story_translation.language in languages_currently_translating:
                self.logger.error("User can't be assigned as a translator")
                raise Exception("User can't be assigned as a translator")
            else:
                db.session.add(new_story_translation)
                db.session.commit()
        except Exception as error:
            self.logger.error(str(error))
            raise error

        try:
            self._insert_empty_story_translation_contents(db, new_story_translation)

            # Update story.translated_languages
            story = Story.query.filter_by(id=new_story_translation.story_id).first()

            if new_story_translation.language not in story.translated_languages:
                story.translated_languages.append(new_story_translation.language)
                Story.query.filter_by(id=new_story_translation.story_id).update(
                    story.to_dict()
                )

            db.session.commit()
        except Exception as error:
            db.session.delete(new_story_translation)
            db.session.commit()
            self.logger.error(str(error))
            raise error

        return {**story.to_dict(), **new_story_translation.to_dict()}

    def create_translation_test(self, user_id, level, language, wants_reviewer):
        try:
            story_translation_tests = (
                StoryTranslation.query.filter(StoryTranslation.translator_id == user_id)
                .filter(StoryTranslation.language == language)
                .filter(StoryTranslation.is_test == True)
                .all()
            )
            last_30_days = datetime.utcnow() - timedelta(days=30)

            for story_translation_test in story_translation_tests:
                if story_translation_test.stage != "PUBLISH":
                    self.logger.error(
                        f"User has an ongoing story translation test for language {language}."
                    )
                    raise Exception(
                        f"User has an ongoing story translation test for language {language}."
                    )
                elif (
                    story_translation_test.test_result == {}
                    and story_translation_test.reviewer_last_activity
                    and last_30_days < story_translation_test.reviewer_last_activity
                ):
                    self.logger.error("User has failed a test within the last 30 days.")
                    raise Exception("User has failed a test within the last 30 days.")

            test_story = (
                Story.query.filter(Story.is_test == True)
                .filter(Story.level == level)
                .first()
            )
            new_story_translation = StoryTranslation(
                story_id=test_story.id,
                language=language,
                stage="TRANSLATE",
                translator_id=user_id,
                is_test=True,
                test_result={"wants_reviewer": wants_reviewer},
            )
            db.session.add(new_story_translation)
            db.session.commit()
        except Exception as error:
            self.logger.error(str(error))
            raise error
        try:
            self._insert_empty_story_translation_contents(db, new_story_translation)
        except Exception as error:
            db.session.delete(new_story_translation)
            db.session.commit()
            self.logger.error(str(error))
            raise error
        return new_story_translation

    def get_story_translations_by_user(
        self, user_id, is_translator=None, language=None, level=None
    ):
        if is_translator is None:
            role_filter = (StoryTranslationAll.translator_id == user_id) | (
                StoryTranslationAll.reviewer_id == user_id
            )
        elif is_translator:
            role_filter = StoryTranslationAll.translator_id == user_id
        else:
            role_filter = StoryTranslationAll.reviewer_id == user_id
        return self.get_story_translations(
            language=language, level=level, role_filter=role_filter
        )

    def get_story_translations(
        self,
        language=None,
        level=None,
        stage=None,
        story_title=None,
        story_id=None,
        role_filter=None,
    ):
        try:
            filters = [
                StoryTranslation.is_test == False,
                StoryTranslationAll.is_test == False,
            ]
            if language is not None:
                filters.append(StoryTranslation.language == language)
                filters.append(StoryTranslationAll.language == language)
            if stage is not None:
                filters.append(StoryTranslation.stage == stage)
                filters.append(StoryTranslationAll.stage == stage)
            if level is not None:
                filters.append(Story.level == level)
            if story_title is not None:
                filters.append(Story.title.like(f"%{story_title}%"))
            if role_filter is not None:
                filters.append(role_filter)
            if story_id is not None:
                filters.append(StoryTranslation.story_id == story_id)
                filters.append(StoryTranslationAll.story_id == story_id)

            stories = (
                Story.query.join(
                    StoryTranslation,
                    Story.id == StoryTranslation.story_id,
                )
                .filter(*filters)
                .order_by(Story.id)
                .all()
            )

            translator = aliased(User)
            reviewer = aliased(User)

            story_translation_data = (
                db.session.query(
                    StoryTranslationAll,
                    translator,
                    reviewer,
                )
                .join(Story, Story.id == StoryTranslationAll.story_id)
                .join(
                    translator,
                    StoryTranslationAll.translator_id == translator.id,
                    isouter=True,
                )
                .join(
                    reviewer,
                    StoryTranslationAll.reviewer_id == reviewer.id,
                    isouter=True,
                )
                .join("translation_contents")
                .filter(*filters)
                .filter(StoryTranslationAll.is_deleted == False)
                .order_by(Story.id)
                .all()
            )

            story_translations = []

            for i in range(len(story_translation_data)):
                story_translation, translator, reviewer = story_translation_data[i]
                story = next(
                    filter(lambda x: x.id == story_translation.story_id, stories)
                )
                translation_contents = list(
                    map(lambda x: x.to_dict(), story_translation.translation_contents)
                )

                story_translations.append(
                    {
                        **story.to_dict(),
                        **story_translation.to_dict(include_relationships=True),
                        "story_translation_id": story_translation.id,
                        "translator_name": (
                            f"{translator.first_name} {translator.last_name}"
                            if translator
                            else None
                        ),
                        "reviewer_name": (
                            f"{reviewer.first_name} {reviewer.last_name}"
                            if reviewer
                            else None
                        ),
                        "num_translated_lines": self._get_num_translated_lines(
                            translation_contents
                        ),
                        "num_approved_lines": self._get_num_approved_lines(
                            translation_contents
                        ),
                        "num_content_lines": len(translation_contents),
                    }
                )
            return story_translations

        except Exception as error:
            self.logger.error(str(error))
            raise error

    def get_story_translation_tests(
        self,
        user,
        language=None,
        level=None,
        stage=None,
        story_title=None,
        submitted_only=False,
    ):
        try:
            filters = [StoryTranslation.is_test]
            if submitted_only:
                filters.append(
                    or_(
                        StoryTranslation.stage == "REVIEW",
                        StoryTranslation.stage == "PUBLISH",
                    )
                )
            if language is not None:
                filters.append(StoryTranslation.language == language)
            if stage is not None:
                filters.append(StoryTranslation.stage == stage)
            if level is not None:
                filters.append(Story.level == level)
            if story_title is not None:
                filters.append(Story.title.like(f"%{story_title}%"))
            if user.role != "Admin":
                filters.append(StoryTranslation.translator_id == user.id)

            stories = (
                Story.query.join(
                    StoryTranslation,
                    Story.id == StoryTranslation.story_id,
                )
                .filter(*filters)
                .order_by(Story.id)
                .all()
            )

            translator = aliased(User)

            story_translation_data = (
                db.session.query(StoryTranslation, translator)
                .join(Story, Story.id == StoryTranslation.story_id)
                .join(
                    translator,
                    StoryTranslation.translator_id == translator.id,
                    isouter=True,
                )
                .filter(*filters)
                .order_by(Story.id)
                .all()
            )

            story_translations = []

            for i in range(len(story_translation_data)):
                story_translation, translator = story_translation_data[i]
                story = next(
                    filter(lambda x: x.id == story_translation.story_id, stories)
                )

                story_translations.append(
                    {
                        **story.to_dict(),
                        **story_translation.to_dict(include_relationships=True),
                        "translator_name": (
                            f"{translator.first_name} {translator.last_name}"
                            if translator
                            else None
                        ),
                    }
                )
            return story_translations

        except Exception as error:
            self.logger.error(str(error))
            raise error

    def get_story_translation(self, id):
        try:
            translator = aliased(User)
            reviewer = aliased(User)

            story_details = (
                db.session.query(
                    Story.id.label("story_id"),
                    Story.title.label("title"),
                    Story.description.label("description"),
                    Story.youtube_link.label("youtube_link"),
                    Story.level.label("level"),
                    Story.is_test.label("is_test"),
                    StoryTranslation.id.label("id"),
                    StoryTranslation.language.label("language"),
                    StoryTranslation.stage.label("stage"),
                    StoryTranslation.translator_id.label("translator_id"),
                    StoryTranslation.reviewer_id.label("reviewer_id"),
                    StoryTranslation.test_result.label("test_result"),
                    StoryTranslation.test_feedback.label("test_feedback"),
                    StoryTranslationContent.id.label("content_id"),
                    StoryTranslationContent.line_index.label("line_index"),
                    StoryTranslationContent.status.label("status"),
                    StoryTranslationContent.translation_content.label(
                        "translation_content"
                    ),
                    (translator.first_name + " " + translator.last_name).label(
                        "translator_name"
                    ),
                    (reviewer.first_name + " " + reviewer.last_name).label(
                        "reviewer_name"
                    ),
                )
                .join(StoryTranslation, Story.id == StoryTranslation.story_id)
                .join(
                    StoryTranslationContent,
                    StoryTranslationContent.story_translation_id == StoryTranslation.id,
                )
                .join(
                    translator,
                    StoryTranslation.translator_id == translator.id,
                    isouter=True,
                )
                .join(
                    reviewer,
                    StoryTranslation.reviewer_id == reviewer.id,
                    isouter=True,
                )
                .filter(StoryTranslation.id == id)
                .all()
            )

            response = {**story_details[0]._asdict()}
            response.pop("content_id")
            response.pop("line_index")
            response.pop("status")
            response.pop("translation_content")
            response["translation_contents"] = []

            for story_detail in story_details:
                story_detail_dict = story_detail._asdict()
                response["translation_contents"].append(
                    {
                        "id": story_detail_dict["content_id"],
                        "line_index": story_detail_dict["line_index"],
                        "status": story_detail_dict["status"],
                        "translation_content": story_detail_dict["translation_content"],
                    }
                )

            response["num_translated_lines"] = self._get_num_translated_lines(
                response["translation_contents"]
            )

            response["num_approved_lines"] = self._get_num_approved_lines(
                response["translation_contents"]
            )

            response["num_content_lines"] = len(response["translation_contents"])

            return response

        except Exception as error:
            self.logger.error(str(error))
            raise error

    def assign_user_as_reviewer(self, user, story_translation):
        story_translations_reviewing = (
            self._get_story_translations_user_translating_query(
                user.id, isTranslator=False
            ).all()
        )
        languages_currently_reviewing = self._get_story_translation_languages(
            story_translations_reviewing
        )
        if (
            story_translation["language"] in user.approved_languages_review
            and story_translation["language"] not in languages_currently_reviewing
            and user.approved_languages_review[story_translation["language"]]
            >= story_translation["level"]
            and not story_translation["reviewer_id"]
            and user.id != story_translation["translator_id"]
        ):
            story_translation = StoryTranslation.query.get(story_translation["id"])
            story_translation.reviewer_id = user.id
            db.session.commit()
        else:
            self.logger.error("User can't be assigned as a reviewer")
            raise Exception("User can't be assigned as a reviewer")

        return {
            **self.get_story(story_translation.story_id),
            **story_translation.to_dict(),
        }

    def remove_reviewer_from_story_translation(self, story_translation):
        try:
            story_translation = StoryTranslation.query.get(story_translation["id"])
            story_translation.reviewer_id = None
            db.session.commit()
        except Exception as error:
            reason = getattr(error, "message", None)
            self.logger.error(
                "Failed to remove reviewer from story translation. Reason = {reason}".format(
                    reason=(reason if reason else str(error))
                )
            )
            raise error

    def update_story(self, story_id, title, description, youtube_link):
        try:
            story = Story.query.get(story_id)
            story.title = title
            story.description = description
            story.youtube_link = youtube_link
            db.session.commit()
        except Exception as error:
            reason = getattr(error, "message", None)
            self.logger.error(
                "Failed to update story. Reason = {reason}".format(
                    reason=(reason if reason else str(error))
                )
            )
            raise error

    # Deprecated: function is not currently in use (story translation stage logic has not been tested)
    def update_story_translation_content(self, story_translation_content):
        story_translation = StoryTranslationContent.query.filter_by(
            id=story_translation_content.id
        ).first()
        story_translation_id = story_translation.story_translation_id

        story_translation = StoryTranslation.query.filter_by(
            id=story_translation_id
        ).first()
        story_translation_stage = story_translation.stage

        if story_translation_stage == "TRANSLATE":
            try:
                if not story_translation:
                    raise Exception(
                        "story_translation_content_id {id} not found".format(
                            id=story_translation_content.id
                        )
                    )

                story_translation.translation_content = (
                    story_translation_content.translation_content
                )
                db.session.commit()
            except Exception as error:
                reason = getattr(error, "message", None)
                self.logger.error(
                    "Failed to update story translation content. Reason = {reason}".format(
                        reason=(reason if reason else str(error))
                    )
                )
                raise error

            return StoryTranslationContentResponseDTO(
                story_translation_content.id,
                story_translation.line_index,
                story_translation_content.translation_content,
            )
        elif story_translation_stage == "REVIEW":
            raise Exception(
                "Story translation contents cannot be changed while the story is being reviewed."
            )
        elif story_translation_stage == "PUBLISH":
            raise Exception(
                "Story translation contents cannot be changed after the story has been published."
            )
        else:
            raise Exception("Story translation contents cannot be changed right now.")

    def update_story_translation_contents(self, story_translation_contents):
        try:
            story_translation = (
                StoryTranslation.query.join(
                    StoryTranslationContent,
                    StoryTranslation.id == StoryTranslationContent.story_translation_id,
                )
                .filter(StoryTranslationContent.id == story_translation_contents[0].id)
                .first()
            )

            story_translation_stage = story_translation.stage
        except Exception as error:
            self.logger.error(str(error))
            raise error

        if story_translation_stage == "TRANSLATE":
            try:
                # TODO: return lineIndex too
                db.session.bulk_update_mappings(
                    StoryTranslationContent, story_translation_contents
                )
                db.session.commit()

                self._update_story_translation_last_activity(story_translation, True)

                return story_translation_contents
            except Exception as error:
                reason = getattr(error, "message", None)
                self.logger.error(
                    "Failed to update story translation content. Reason = {reason}".format(
                        reason=(reason if reason else str(error))
                    )
                )
                raise error
        elif story_translation_stage == "REVIEW":
            raise Exception(
                "Story translation contents cannot be changed while the story is being reviewed."
            )
        elif story_translation_stage == "PUBLISH":
            raise Exception(
                "Story translation contents cannot be changed after the story has been published."
            )
        else:
            raise Exception("Story translation contents cannot be changed right now.")

    def update_story_translation_stage(self, story_translation_data, user_id):
        try:
            story_translation = StoryTranslation.query.filter_by(
                id=story_translation_data["id"]
            ).first()
            new_stage = story_translation_data["stage"]

            is_translator = user_id == story_translation.translator_id
            is_reviewer = user_id == story_translation.reviewer_id

            if (
                (new_stage == StageEnum.TRANSLATE or new_stage == StageEnum.PUBLISH)
                and is_reviewer
            ) or (new_stage == StageEnum.REVIEW and is_translator):
                story_translation.stage = new_stage
                db.session.commit()

                self._update_story_translation_last_activity(
                    story_translation, is_translator
                )
            else:
                error = "User is not authorized to update translation stage to: {stage}".format(
                    stage=new_stage
                )
                raise Exception(error)
        except Exception as error:
            self.logger.error(error)
            raise error

    def get_stories_available_for_translation(self, language, level, user_id):
        try:
            ongoing_translations = (
                self._get_story_translations_user_translating_query(
                    user_id, isTranslator=True
                )
                .filter(StoryTranslation.language == language)
                .all()
            )
            if len(ongoing_translations) > 0:
                return []

            stories = (
                Story.query.filter(Story.level == level)
                .filter(~Story.translated_languages.contains(language))
                .all()
            )
            return [story.to_dict(include_relationships=True) for story in stories]
        except Exception as error:
            self.logger.error(str(error))
            raise error

    def get_story_translations_available_for_review(self, language, level, user_id):
        try:
            ongoing_translations = (
                self._get_story_translations_user_translating_query(
                    user_id, isTranslator=False
                )
                .filter(StoryTranslation.language == language)
                .all()
            )
            if len(ongoing_translations) > 0:
                return []

            stories = (
                Story.query.join(
                    StoryTranslation, Story.id == StoryTranslation.story_id
                )
                .filter(Story.level == level)
                .filter(StoryTranslation.language == language)
                .filter(StoryTranslation.reviewer_id == None)
                .order_by(Story.id)
                .all()
            )

            story_translations = (
                StoryTranslation.query.join(
                    Story, Story.id == StoryTranslation.story_id
                )
                .filter(Story.level == level)
                .filter(StoryTranslation.language == language)
                .filter(StoryTranslation.reviewer_id == None)
                .order_by(Story.id)
                .all()
            )

            for i in range(len(stories)):
                stories[i] = {
                    **stories[i].to_dict(),
                    **story_translations[i].to_dict(include_relationships=True),
                }

            return stories

        except Exception as error:
            self.logger.error(str(error))
            raise error

    def update_story_translation_content_status(
        self, story_translation_content_id, status
    ):
        try:
            story_translation_content = StoryTranslationContent.query.filter_by(
                id=story_translation_content_id
            ).first()

            story_translation = StoryTranslation.query.filter_by(
                id=story_translation_content.story_translation_id
            ).first()

            if story_translation.stage == "REVIEW":
                if story_translation.is_test:
                    if not "TEST_" in status and status != "DEFAULT":
                        raise Exception(
                            "Error. Story Translation Test cannot have non-test statuses on story translation content."
                        )
                else:
                    if "TEST_" in status:
                        raise Exception(
                            "Error. Story Translation cannot have test statuses on story translation content."
                        )

                story_translation_content.status = status
                db.session.commit()

                self._update_story_translation_last_activity(story_translation, False)
            else:
                raise Exception(
                    "Error. Story Translation is not in REVIEW stage and its statuses cannot be updated."
                )
        except Exception as error:
            reason = getattr(error, "message", None)
            self.logger.error(
                "Failed to update story translation status. Reason = {reason}".format(
                    reason=(reason if reason else str(error))
                )
            )
            raise error

        return StoryTranslationUpdateStatusResponseDTO(
            story_translation_content_id,
            story_translation_content.line_index,
            status,
        )

    def approve_all_story_translation_content(self, story_translation_id):
        try:
            story_translation = StoryTranslation.query.get(story_translation_id)
            story_translation_contents = (
                db.session.query(StoryTranslationContent)
                .select_from(StoryTranslation)
                .join(
                    StoryTranslationContent,
                    StoryTranslation.id == StoryTranslationContent.story_translation_id,
                )
                .filter(StoryTranslation.id == story_translation_id)
                .all()
            )

            if story_translation.stage == "REVIEW":
                for translation_content in story_translation_contents:
                    translation_content.status = "APPROVED"

                db.session.commit()
                self._update_story_translation_last_activity(story_translation, False)
            else:
                raise Exception(
                    "Error. Story Translation is not in REVIEW stage and its statuses cannot be updated."
                )
        except Exception as error:
            reason = getattr(error, "message", None)
            self.logger.error(
                "Failed to update all story translation status. Reason = {reason}".format(
                    reason=(reason if reason else str(error))
                )
            )
            raise error

    def finish_grading_story_translation(
        self, reviewer_id, test_result, test_feedback, story_translation_test_id
    ):
        try:
            story_translation_test = (
                db.session.query(StoryTranslation)
                .filter(StoryTranslation.id == story_translation_test_id)
                .first()
            )

            if story_translation_test.is_test == False:
                raise Exception("Error: Cannot grade a non-test story translation.")

            story_translation_contents = (
                db.session.query(StoryTranslationContent)
                .filter(
                    StoryTranslationContent.story_translation_id
                    == story_translation_test_id
                )
                .all()
            )

            language = story_translation_test.language
            user = (
                db.session.query(User)
                .filter(User.id == story_translation_test.translator_id)
                .first()
            )

            if "translate" in test_result:
                translate_level = test_result["translate"]
                appr_lang = user.approved_languages_translation or {}
                appr_lang[language] = translate_level
                user.approved_languages_translation = appr_lang
            else:
                test_result = {}

            if "review" in test_result:
                review_level = test_result["review"]
                appr_lang = user.approved_languages_review or {}
                appr_lang[language] = review_level
                user.approved_languages_review = appr_lang

            User.query.filter_by(id=user.id).update(user.to_dict())

            story_translation_test.reviewer_id = reviewer_id
            story_translation_test.test_feedback = test_feedback
            story_translation_test.test_result = test_result
            story_translation_test.test_grade = self._get_final_grade(
                story_translation_contents
            )
            story_translation_test.stage = "PUBLISH"
            db.session.commit()
            self._update_story_translation_last_activity(story_translation_test, False)

        except Exception as error:
            reason = getattr(error, "message", None)
            self.logger.error(
                "Failed to update story translation. Reason = {reason}".format(
                    reason=(reason if reason else str(error))
                )
            )
            raise error

        return story_translation_test

    def soft_delete_story_translation(self, id):
        try:
            story_translation = StoryTranslationAll.query.get(id)
            for translation_content in story_translation.translation_contents:
                comments = Comment.query.filter(
                    Comment.story_translation_content_id == translation_content.id
                )
                for comment in comments:
                    comment.is_deleted = True
                translation_content.is_deleted = True
            story_translation.is_deleted = True

            # Remove language from story.translated_languages
            story = Story.query.filter_by(id=story_translation.story_id).first()
            if (
                story is not None
                and story.translated_languages is not None
                and story_translation.language in story.translated_languages
            ):
                story.translated_languages.remove(story_translation.language)
                Story.query.filter_by(id=story_translation.story_id).update(
                    story.to_dict()
                )

            db.session.commit()
        except Exception as error:
            self.logger.error(error)
            raise error

    def remove_user_from_story_translation(self, story_translation_id, user_id):
        try:
            story_translation = StoryTranslation.query.get(story_translation_id)
            if not story_translation:
                raise Exception("Error. Story translation does not exist.")
            if user_id == story_translation.translator_id:
                self.soft_delete_story_translation(story_translation_id)
            elif user_id == story_translation.reviewer_id:
                story_translation.reviewer_id = None
                db.session.commit()
            else:
                raise Exception(
                    "Error. User is not a translator or reviewer of this story translation."
                )
        except Exception as error:
            self.logger.error(error)
            raise error

    def get_story_translation_statistics(self):
        num_translations_in_translation = (
            db.session.query(StoryTranslation)
            .filter(StoryTranslation.stage == "TRANSLATE")
            .filter(StoryTranslation.is_test == False)
            .count()
        )
        num_translations_in_review = (
            db.session.query(StoryTranslation)
            .filter(StoryTranslation.stage == "REVIEW")
            .filter(StoryTranslation.is_test == False)
            .count()
        )
        num_translations_completed = (
            db.session.query(StoryTranslation)
            .filter(StoryTranslation.stage == "PUBLISH")
            .filter(StoryTranslation.is_test == False)
            .count()
        )
        return StoryTranslationStatisticsResponseDTO(
            num_translations_in_translation=num_translations_in_translation,
            num_translations_in_review=num_translations_in_review,
            num_translations_completed=num_translations_completed,
        )

    def soft_delete_story(self, id):
        try:
            story = StoryAll.query.get(id)
            for content in story.contents:
                content.is_deleted = True

            story.is_deleted = True

            # Remove story_translations for this story
            story_translations = (
                db.session.query(StoryTranslation)
                .filter(StoryTranslation.story_id == id)
                .all()
            )

            for translation in story_translations:
                self.soft_delete_story_translation(translation.id)

            db.session.commit()
        except Exception as error:
            self.logger.error(error)
            raise error

    def _get_num_translated_lines(self, translation_contents):
        return len(translation_contents) - [
            _["translation_content"].strip() for _ in translation_contents
        ].count("")

    def _get_num_approved_lines(self, translation_contents):
        count = 0
        for translation_content in translation_contents:
            if translation_content["status"] == StoryTranslationContentStatus.APPROVED:
                count += 1

        return count

    def _get_story_translations_user_translating_query(self, user_id, isTranslator):
        return (
            StoryTranslation.query.filter(StoryTranslation.is_test == False)
            .filter(
                StoryTranslation.translator_id == user_id
                if isTranslator
                else StoryTranslation.reviewer_id == user_id
            )
            .filter(StoryTranslation.stage != "PUBLISH")
        )

    def _get_story_translation_languages(self, story_translations):
        return set(
            [story_translation.language for story_translation in story_translations]
        )

    def _insert_empty_story_translation_contents(self, db, new_story_translation):
        translation_id = new_story_translation.id
        story_id = new_story_translation.story_id
        story_lines = len(StoryContent.query.filter_by(story_id=story_id).all())
        new_story_translation_content_cache = [
            StoryTranslationContent(
                story_translation_id=translation_id,
                line_index=line,
                translation_content="",
            )
            for line in range(story_lines)
        ]
        db.session.bulk_save_objects(new_story_translation_content_cache)
        db.session.commit()

    def _get_final_grade(self, story_translation_contents):

        final_grade = 0
        for s in story_translation_contents:
            status = s.status
            if status == StoryTranslationContentStatus.TEST_CORRECT:
                final_grade += 1
            elif status == StoryTranslationContentStatus.TEST_PARTIALLY_CORRECT:
                final_grade += 0.5

        return final_grade

    def _update_story_translation_last_activity(self, story_translation, is_translator):
        try:
            if not story_translation:
                raise Exception("Error. Story translation does not exist.")
            if is_translator:
                story_translation.translator_last_activity = datetime.utcnow()
                db.session.commit()
            else:
                story_translation.reviewer_last_activity = datetime.utcnow()
                db.session.commit()
        except Exception as error:
            self.logger.error(error)
            raise error

    def process_story(self, file):
        try:
            doc = docx.Document(file)
            story_contents = []
            for line in doc.paragraphs:
                story_contents.append(line.text)
            return story_contents
        except Exception as e:
            error_message = getattr(e, "message", None)
            raise Exception(error_message if error_message else str(e))
